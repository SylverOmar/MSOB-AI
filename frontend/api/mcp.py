"""Stateless Streamable HTTP MCP endpoint for MSOB clinical documents.

The browser places temporary analysis files in a private Supabase queue. The
Agentic workflow calls the single ``ocr_extract_document`` tool with the queue
request ID. This function claims the queue, processes each file sequentially,
and asks the Supabase broker to delete every temporary object afterward.
"""

from __future__ import annotations

import base64
import io
import json
import mimetypes
import tempfile
import urllib.error
import urllib.parse
import urllib.request
import zipfile
from pathlib import Path
from typing import Any

from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse, Response


SUPABASE_URL = "https://inqnyoqqhtogrvgjyavo.supabase.co"
PUBLISHABLE_KEY = "sb_publishable_m45iFH-I3oUdBtz8Bl0_-g_DemQoMoH"
PUBLIC_CONFIG_URL = f"{SUPABASE_URL}/rest/v1/rpc/msob_public_runtime_config"
DEFAULT_BROKER_URL = f"{SUPABASE_URL}/functions/v1/msob-medical-files"
MAX_DOWNLOAD_BYTES = 25 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp"}
ALLOWED_ORIGINS = {
    "https://msob-ai.vercel.app",
    "https://stg-agentic.abafusion.ai",
}

app = FastAPI(title="MSOB Clinical Documents MCP", docs_url=None, redoc_url=None)


class McpProcessingError(RuntimeError):
    pass


def _request_json(url: str, payload: dict[str, Any], headers: dict[str, str] | None = None) -> Any:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=body,
        method="POST",
        headers={"Content-Type": "application/json", **(headers or {})},
    )
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            raw = response.read().decode("utf-8", errors="replace")
            return json.loads(raw) if raw else {}
    except urllib.error.HTTPError as error:
        raw = error.read().decode("utf-8", errors="replace")
        try:
            detail = json.loads(raw).get("error") or json.loads(raw).get("message")
        except (json.JSONDecodeError, AttributeError):
            detail = raw
        raise McpProcessingError(str(detail or f"Service error {error.code}")) from error
    except urllib.error.URLError as error:
        raise McpProcessingError("The temporary document service is unavailable.") from error


def _runtime_config() -> dict[str, Any]:
    try:
        value = _request_json(
            PUBLIC_CONFIG_URL,
            {},
            {
                "apikey": PUBLISHABLE_KEY,
                "Authorization": f"Bearer {PUBLISHABLE_KEY}",
            },
        )
        return value if isinstance(value, dict) else {}
    except McpProcessingError:
        return {}


def _broker_url() -> str:
    configured = str(_runtime_config().get("hosted_mcp_queue_url") or "").strip()
    return configured if configured.startswith("https://") else DEFAULT_BROKER_URL


def _broker(operation: str, **payload: Any) -> dict[str, Any]:
    result = _request_json(_broker_url(), {"operation": operation, **payload})
    if not isinstance(result, dict):
        raise McpProcessingError("The temporary document service returned an invalid response.")
    return result


def _download(url: str) -> tuple[bytes, str]:
    request = urllib.request.Request(url, headers={"User-Agent": "MSOB-AI-MCP/1.0"})
    try:
        with urllib.request.urlopen(request, timeout=60) as response:
            length = int(response.headers.get("Content-Length") or 0)
            if length > MAX_DOWNLOAD_BYTES:
                raise McpProcessingError("A document exceeds the 25 MB limit.")
            content = response.read(MAX_DOWNLOAD_BYTES + 1)
            if len(content) > MAX_DOWNLOAD_BYTES:
                raise McpProcessingError("A document exceeds the 25 MB limit.")
            return content, response.headers.get_content_type() or "application/octet-stream"
    except urllib.error.URLError as error:
        raise McpProcessingError("A temporary document could not be downloaded.") from error


def _normalized_image(image_bytes: bytes, mime_type: str) -> tuple[bytes, str]:
    try:
        from PIL import Image, ImageOps

        with Image.open(io.BytesIO(image_bytes)) as image:
            image = ImageOps.exif_transpose(image).convert("RGB")
            image.thumbnail((2200, 2200))
            output = io.BytesIO()
            image.save(output, format="JPEG", quality=86, optimize=True)
            return output.getvalue(), "image/jpeg"
    except Exception:
        return image_bytes, mime_type


def _vision_text(
    request_id: str,
    claim_token: str,
    image_bytes: bytes,
    mime_type: str,
    prompt: str,
) -> str:
    normalized, normalized_mime = _normalized_image(image_bytes, mime_type)
    result = _broker(
        "mcp-vision",
        request_id=request_id,
        claim_token=claim_token,
        mime_type=normalized_mime,
        image_base64=base64.b64encode(normalized).decode("ascii"),
        prompt=prompt,
    )
    return str(result.get("text") or "").strip()


def _pdf_parts(path: Path) -> tuple[str, list[tuple[bytes, str]]]:
    text_parts: list[str] = []
    images: list[tuple[bytes, str]] = []
    try:
        from pypdf import PdfReader

        for index, page in enumerate(PdfReader(str(path)).pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(f"Page {index}\n{text.strip()}")
    except Exception as error:
        text_parts.append(f"PDF text extraction failed: {error}")

    try:
        import fitz

        document = fitz.open(str(path))
        for page in document:
            for image_info in page.get_images(full=True):
                image = document.extract_image(image_info[0])
                extension = str(image.get("ext") or "png").lower()
                mime = "image/jpeg" if extension in {"jpg", "jpeg"} else f"image/{extension}"
                images.append((image["image"], mime))
        document.close()
    except Exception:
        pass
    return "\n\n".join(text_parts), images


def _docx_parts(path: Path) -> tuple[str, list[tuple[bytes, str]]]:
    text_parts: list[str] = []
    images: list[tuple[bytes, str]] = []
    try:
        import docx

        document = docx.Document(str(path))
        text_parts.extend(p.text.strip() for p in document.paragraphs if p.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(dict.fromkeys(cells)))
    except Exception as error:
        text_parts.append(f"DOCX text extraction failed: {error}")

    try:
        with zipfile.ZipFile(path) as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                extension = name.rsplit(".", 1)[-1].lower()
                mime = "image/jpeg" if extension in {"jpg", "jpeg"} else f"image/{extension}"
                images.append((archive.read(name), mime))
    except Exception:
        pass
    return "\n".join(text_parts), images


def _extract_file(
    path: Path,
    request_id: str,
    claim_token: str,
    prompt: str,
    content_type: str,
) -> str:
    extension = path.suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        return f"Unsupported file format: {extension or '(none)'}"
    if extension in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")

    extracted_text = ""
    images: list[tuple[bytes, str]] = []
    if extension == ".pdf":
        extracted_text, images = _pdf_parts(path)
    elif extension == ".docx":
        extracted_text, images = _docx_parts(path)
    else:
        mime = content_type if content_type.startswith("image/") else mimetypes.guess_type(path.name)[0]
        images = [(path.read_bytes(), mime or "image/png")]

    parts: list[str] = []
    if extracted_text.strip():
        parts.append("=== EXTRACTED TEXT ===\n" + extracted_text.strip())
    for index, (image_bytes, mime_type) in enumerate(images, start=1):
        observation = _vision_text(
            request_id,
            claim_token,
            image_bytes,
            mime_type,
            prompt,
        )
        parts.append(f"=== IMAGE EXTRACTION {index} ({mime_type}) ===\n{observation}")
    return "\n\n".join(parts) or "No usable content could be extracted."


def _process_queue(request_id: str, prompt: str) -> str:
    claim = _broker("mcp-claim", request_id=request_id)
    claim_token = str(claim.get("claim_token") or "")
    files = claim.get("files") if isinstance(claim.get("files"), list) else []
    if not claim_token:
        raise McpProcessingError("The temporary document claim is invalid.")

    results: list[str] = []
    try:
        with tempfile.TemporaryDirectory(prefix="msob-mcp-") as temp_dir:
            root = Path(temp_dir)
            for index, item in enumerate(files, start=1):
                if not isinstance(item, dict):
                    continue
                original_name = Path(str(item.get("name") or f"document-{index}")).name
                extension = Path(original_name).suffix.lower()
                local_path = root / f"{index:03d}{extension}"
                content, response_type = _download(str(item.get("signed_url") or ""))
                local_path.write_bytes(content)
                source = "MEDICAL FOLDER FILE" if item.get("source") == "medical-folder" else "CURRENT CASE FILE"
                result = _extract_file(
                    local_path,
                    request_id,
                    claim_token,
                    prompt,
                    str(item.get("content_type") or response_type),
                )
                results.append(f"=== {source}: {original_name} ===\n{result}")
    finally:
        try:
            _broker("mcp-finalize", request_id=request_id, claim_token=claim_token)
        except Exception:
            pass
    return "\n\n".join(results) or f"No files are queued for clinical request {request_id}."


def _direct_input(input_value: str) -> str:
    value = str(input_value or "").strip()
    if not value:
        raise McpProcessingError("Provide request_id for queued files or input_value for plain text.")
    parsed = urllib.parse.urlparse(value)
    if parsed.scheme in {"http", "https"}:
        raise McpProcessingError(
            "Direct document URLs are not accepted. Use request_id for queued files."
        )
    return value


def _tool_call(arguments: dict[str, Any]) -> str:
    request_id = str(arguments.get("request_id") or "").strip()
    prompt = str(arguments.get("prompt") or "").strip() or (
        "Extract the document's clinical text and observations accurately. "
        "Do not diagnose, infer missing facts, or add recommendations."
    )
    if request_id:
        return _process_queue(request_id, prompt)
    return _direct_input(str(arguments.get("input_value") or ""))


def _jsonrpc_result(request_id: Any, result: dict[str, Any]) -> JSONResponse:
    return JSONResponse({"jsonrpc": "2.0", "id": request_id, "result": result})


def _jsonrpc_error(request_id: Any, code: int, message: str) -> JSONResponse:
    return JSONResponse({
        "jsonrpc": "2.0",
        "id": request_id,
        "error": {"code": code, "message": message},
    })


@app.get("/")
@app.get("/api/mcp")
async def mcp_info() -> JSONResponse:
    return JSONResponse({
        "status": "ok",
        "service": "MSOB Clinical Documents MCP",
        "transport": "streamable-http",
        "tool": "ocr_extract_document",
    })


@app.delete("/")
@app.delete("/api/mcp")
async def close_stateless_session() -> Response:
    return Response(status_code=204)


@app.post("/")
@app.post("/api/mcp")
async def mcp_endpoint(request: Request) -> Response:
    origin = request.headers.get("origin")
    if origin and origin not in ALLOWED_ORIGINS and not origin.endswith(".vercel.app"):
        return _jsonrpc_error(None, -32000, "Origin not allowed.")
    try:
        payload = await request.json()
    except Exception:
        return _jsonrpc_error(None, -32700, "Parse error")
    if not isinstance(payload, dict):
        return _jsonrpc_error(None, -32600, "Invalid Request")

    method = str(payload.get("method") or "")
    request_id = payload.get("id")
    params = payload.get("params") if isinstance(payload.get("params"), dict) else {}

    if method.startswith("notifications/"):
        return Response(status_code=202)
    if method == "initialize":
        offered = str(params.get("protocolVersion") or "2025-03-26")
        return _jsonrpc_result(request_id, {
            "protocolVersion": offered,
            "capabilities": {"tools": {"listChanged": False}},
            "serverInfo": {"name": "MSOB-Clinical-Documents", "version": "1.0.0"},
        })
    if method == "ping":
        return _jsonrpc_result(request_id, {})
    if method == "tools/list":
        return _jsonrpc_result(request_id, {"tools": [{
            "name": "ocr_extract_document",
            "description": (
                "Extract queued clinical-analysis files one by one. Use request_id for files "
                "prepared by the MSOB website; input_value is only for direct plain text."
            ),
            "inputSchema": {
                "type": "object",
                "properties": {
                    "request_id": {"type": "string"},
                    "input_value": {"type": "string"},
                    "prompt": {"type": "string"},
                },
                "additionalProperties": False,
            },
        }]})
    if method == "tools/call":
        if params.get("name") != "ocr_extract_document":
            return _jsonrpc_error(request_id, -32602, "Unknown tool")
        arguments = params.get("arguments") if isinstance(params.get("arguments"), dict) else {}
        try:
            text = _tool_call(arguments)
            return _jsonrpc_result(request_id, {
                "content": [{"type": "text", "text": text}],
                "isError": False,
            })
        except Exception as error:
            return _jsonrpc_result(request_id, {
                "content": [{"type": "text", "text": str(error)}],
                "isError": True,
            })
    return _jsonrpc_error(request_id, -32601, "Method not found")
