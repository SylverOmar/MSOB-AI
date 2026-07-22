"""MSOB clinical-document MCP server.

The browser uploads the documents needed for one clinical analysis to an
isolated, temporary queue. The Agentic workflow later calls the single
``ocr_extract_document(request_id=...)`` tool to extract the queued documents.
The same tool can still process one direct document through ``input_value``.
Queued files are deleted after processing and are never used for patient
creation or medical-folder persistence.
"""

from __future__ import annotations

import base64
import json
import logging
import mimetypes
import os
import re
import shutil
import time
import uuid
import zipfile
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP

load_dotenv()

ROOT = Path(__file__).resolve().parent
TEMP_HOLD_ROOT = ROOT / "temp_hold"
MANIFEST_NAME = "_manifest.json"
MAX_FILE_BYTES = int(os.getenv("MCP_MAX_FILE_BYTES", str(25 * 1024 * 1024)))
QUEUE_TTL_SECONDS = int(os.getenv("MCP_QUEUE_TTL_SECONDS", str(6 * 60 * 60)))
EXTRACTABLE_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}
ALLOWED_SOURCES = {"case", "medical-folder"}
DEFAULT_BROWSER_ORIGINS = [
    "http://localhost:5500",
    "http://127.0.0.1:5500",
    "https://msob-ai.vercel.app",
]
BROWSER_ORIGINS = sorted(
    {
        *DEFAULT_BROWSER_ORIGINS,
        *(
            origin.strip()
            for origin in os.getenv("MCP_BROWSER_ORIGINS", "").split(",")
            if origin.strip()
        ),
    }
)

mcp = FastMCP(
    "MSOB-Clinical-Documents",
    host=os.getenv("MCP_HOST", "0.0.0.0"),
    port=int(os.getenv("MCP_PORT", "8002")),
    json_response=True,
)

logging.basicConfig(
    level=os.getenv("MCP_LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
)
logger = logging.getLogger("mcp_ocr_server")


def _safe_request_id(value: str) -> str:
    request_id = str(value or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]{8,80}", request_id):
        raise ValueError("Invalid request_id.")
    return request_id


def _safe_original_name(value: str) -> str:
    name = Path(str(value or "document")).name.strip()
    if not name:
        name = "document"
    name = re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", "_", name)
    return name[:180]


def _queue_dir(request_id: str) -> Path:
    return TEMP_HOLD_ROOT / _safe_request_id(request_id)


def _manifest_path(queue_dir: Path) -> Path:
    return queue_dir / MANIFEST_NAME


def _read_manifest(queue_dir: Path) -> list[dict[str, Any]]:
    path = _manifest_path(queue_dir)
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (OSError, json.JSONDecodeError):
        logger.warning("Ignoring an unreadable queue manifest at %s", path)
        return []


def _write_manifest(queue_dir: Path, entries: list[dict[str, Any]]) -> None:
    queue_dir.mkdir(parents=True, exist_ok=True)
    path = _manifest_path(queue_dir)
    temporary = path.with_suffix(".tmp")
    temporary.write_text(
        json.dumps(entries, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    os.replace(temporary, path)


def _cleanup_stale_queues() -> None:
    if not TEMP_HOLD_ROOT.exists():
        return
    now = time.time()
    for queue_dir in TEMP_HOLD_ROOT.iterdir():
        if not queue_dir.is_dir():
            continue
        try:
            if now - queue_dir.stat().st_mtime > QUEUE_TTL_SECONDS:
                shutil.rmtree(queue_dir)
                logger.info("Removed expired clinical queue %s", queue_dir.name)
        except OSError as error:
            logger.warning("Could not inspect stale queue %s: %s", queue_dir, error)


def _clear_all_queues() -> dict[str, int]:
    """Atomically detach and delete every held clinical-document queue."""

    TEMP_HOLD_ROOT.parent.mkdir(parents=True, exist_ok=True)
    TEMP_HOLD_ROOT.mkdir(parents=True, exist_ok=True)
    quarantine = TEMP_HOLD_ROOT.with_name(
        f"{TEMP_HOLD_ROOT.name}.clearing-{uuid.uuid4().hex}"
    )
    try:
        os.replace(TEMP_HOLD_ROOT, quarantine)
    except FileNotFoundError:
        TEMP_HOLD_ROOT.mkdir(parents=True, exist_ok=True)
        return {"requests": 0, "files": 0}

    TEMP_HOLD_ROOT.mkdir(parents=True, exist_ok=True)
    request_count = sum(1 for item in quarantine.iterdir() if item.is_dir())
    file_count = sum(1 for item in quarantine.rglob("*") if item.is_file())
    shutil.rmtree(quarantine, ignore_errors=True)
    logger.info(
        "Cleared all clinical queues (%s requests, %s temporary files)",
        request_count,
        file_count,
    )
    return {"requests": request_count, "files": file_count}


def _store_queued_file(
    request_id: str,
    original_name: str,
    content: bytes,
    source: str,
    content_type: str | None,
) -> dict[str, Any]:
    request_id = _safe_request_id(request_id)
    source = str(source or "case").strip().lower()
    if source not in ALLOWED_SOURCES:
        raise ValueError("Invalid document source.")
    if not content:
        raise ValueError("The uploaded file is empty.")
    if len(content) > MAX_FILE_BYTES:
        raise ValueError(f"The uploaded file exceeds {MAX_FILE_BYTES // (1024 * 1024)} MB.")

    original_name = _safe_original_name(original_name)
    extension = Path(original_name).suffix.lower()
    if not re.fullmatch(r"\.[a-z0-9]{1,12}", extension):
        extension = ".bin"

    queue_dir = _queue_dir(request_id)
    queue_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid.uuid4().hex}{extension}"
    stored_path = queue_dir / stored_name
    stored_path.write_bytes(content)

    entry = {
        "stored_name": stored_name,
        "original_name": original_name,
        "source": source,
        "content_type": content_type
        or mimetypes.guess_type(original_name)[0]
        or "application/octet-stream",
        "size": len(content),
        "uploaded_at": int(time.time()),
    }
    entries = _read_manifest(queue_dir)
    entries.append(entry)
    _write_manifest(queue_dir, entries)
    logger.info(
        "Queued %s (%s bytes) for request %s from %s",
        original_name,
        len(content),
        request_id,
        source,
    )
    return entry


def get_groq_client(use_fallback: bool = False):
    from groq import Groq

    key_name = "GROQ_FALLBACK_API_KEY" if use_fallback else "GROQ_API_KEY"
    api_key = os.getenv(key_name, "") or os.getenv("GROQ_API_KEY", "")
    if not api_key:
        raise RuntimeError("No Groq API key is configured for the MCP server.")
    return Groq(api_key=api_key)


def analyze_image_bytes(image_bytes: bytes, mime_type: str, prompt_text: str) -> str:
    encoded = base64.b64encode(image_bytes).decode("ascii")
    model_name = os.getenv(
        "GROQ_VISION_MODEL",
        "meta-llama/llama-4-scout-17b-16e-instruct",
    )
    last_error: Exception | None = None

    for use_fallback in (False, True):
        try:
            response = get_groq_client(use_fallback).chat.completions.create(
                model=model_name,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt_text},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{encoded}"
                                },
                            },
                        ],
                    }
                ],
                temperature=0.1,
                max_tokens=4096,
            )
            return response.choices[0].message.content or ""
        except Exception as error:  # Provider failures are returned to the workflow.
            last_error = error
            logger.error("Groq image extraction failed: %s", error)

    return f"Image extraction failed: {last_error}"


def extract_from_pdf(filepath: str) -> tuple[str, list[tuple[bytes, str]]]:
    text_parts: list[str] = []
    images: list[tuple[bytes, str]] = []

    try:
        from pypdf import PdfReader

        for index, page in enumerate(PdfReader(filepath).pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {index + 1} ---\n{text}")
    except Exception as error:
        logger.error("PDF text extraction failed: %s", error)
        text_parts.append(f"PDF text extraction failed: {error}")

    try:
        import fitz

        document = fitz.open(filepath)
        for page_index, page in enumerate(document):
            for image_info in page.get_images(full=True):
                try:
                    image = document.extract_image(image_info[0])
                    extension = image["ext"].lower()
                    mime = (
                        "image/jpeg"
                        if extension in {"jpg", "jpeg"}
                        else f"image/{extension}"
                    )
                    images.append((image["image"], mime))
                except Exception as error:
                    logger.warning(
                        "PDF image extraction failed on page %s: %s",
                        page_index + 1,
                        error,
                    )
        document.close()
    except Exception as error:
        logger.error("PDF embedded-image extraction failed: %s", error)

    return "\n\n".join(text_parts), images


def extract_from_docx(filepath: str) -> tuple[str, list[tuple[bytes, str]]]:
    text_parts: list[str] = []
    images: list[tuple[bytes, str]] = []

    try:
        import docx

        document = docx.Document(filepath)
        text_parts.extend(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        for table in document.tables:
            for row in table.rows:
                cells: list[str] = []
                for cell in row.cells:
                    value = cell.text.strip()
                    if value and (not cells or cells[-1] != value):
                        cells.append(value)
                if cells:
                    text_parts.append(" | ".join(cells))
    except Exception as error:
        logger.error("DOCX text extraction failed: %s", error)
        text_parts.append(f"DOCX text extraction failed: {error}")

    try:
        with zipfile.ZipFile(filepath) as archive:
            for name in archive.namelist():
                if not name.startswith("word/media/"):
                    continue
                extension = name.rsplit(".", 1)[-1].lower()
                mime = (
                    "image/jpeg"
                    if extension in {"jpg", "jpeg"}
                    else f"image/{extension}"
                )
                images.append((archive.read(name), mime))
    except Exception as error:
        logger.error("DOCX embedded-image extraction failed: %s", error)

    return "\n".join(text_parts), images


def _extract_document(input_value: str, prompt: str = "") -> str:
    analysis_prompt = str(prompt or "").strip() or (
        "Extract the document's clinical text and observations accurately. "
        "Do not diagnose, infer missing facts, or add recommendations."
    )
    clean_input = str(input_value or "")
    clean_path = clean_input.strip().strip('"').strip("'")
    path = Path(clean_path)

    if not path.is_file():
        return clean_input

    extension = path.suffix.lower()
    extracted_text = ""
    images: list[tuple[bytes, str]] = []

    if extension not in EXTRACTABLE_EXTENSIONS:
        return (
            f"Unsupported file format for extraction: {extension or '(none)'}. "
            "The file was queued successfully but was not executed or interpreted."
        )
    if extension in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if extension in {".png", ".jpg", ".jpeg", ".webp"}:
        mime = "image/jpeg" if extension in {".jpg", ".jpeg"} else f"image/{extension[1:]}"
        images.append((path.read_bytes(), mime))
    elif extension == ".pdf":
        extracted_text, images = extract_from_pdf(str(path))
    elif extension == ".docx":
        extracted_text, images = extract_from_docx(str(path))

    parts: list[str] = []
    if extracted_text.strip():
        parts.append("=== EXTRACTED TEXT ===\n" + extracted_text.strip())
    for index, (image_bytes, mime_type) in enumerate(images, start=1):
        observation = analyze_image_bytes(image_bytes, mime_type, analysis_prompt)
        parts.append(
            f"=== IMAGE EXTRACTION {index} ({mime_type}) ===\n{observation}"
        )
    return "\n\n".join(parts) or "No usable content could be extracted."


def _process_queued_documents(request_id: str, prompt: str = "") -> str:
    try:
        queue_dir = _queue_dir(request_id)
    except ValueError as error:
        return f"Queue error: {error}"

    manifest = _read_manifest(queue_dir)
    if not manifest:
        return f"No files are queued for clinical request {request_id}."

    results: list[str] = []
    for entry in manifest:
        stored_path = queue_dir / str(entry.get("stored_name", ""))
        original_name = str(entry.get("original_name") or stored_path.name)
        source = str(entry.get("source") or "case")
        source_label = (
            "MEDICAL FOLDER FILE" if source == "medical-folder" else "CURRENT CASE FILE"
        )
        try:
            if not stored_path.is_file():
                raise FileNotFoundError("Temporary file is missing.")
            result = _extract_document(str(stored_path), prompt)
            results.append(f"=== {source_label}: {original_name} ===\n{result}")
        except Exception as error:
            logger.exception("Failed to process queued file %s", original_name)
            results.append(
                f"=== {source_label}: {original_name} ===\n"
                f"Document processing failed: {error}"
            )
        finally:
            try:
                stored_path.unlink(missing_ok=True)
            except OSError as error:
                logger.warning("Could not delete temporary file %s: %s", stored_path, error)

    shutil.rmtree(queue_dir, ignore_errors=True)
    logger.info("Processed and cleared clinical queue %s", request_id)
    return "\n\n".join(results)


@mcp.tool()
def ocr_extract_document(
    input_value: str = "",
    request_id: str = "",
    prompt: str = "",
) -> str:
    """Extract one document or a website queue through one MCP tool.

    Use ``request_id`` for a website-launched clinical analysis. The isolated
    queue is processed file by file and deleted afterward. Use ``input_value``
    only for one direct local document or plain-text value. If both are given,
    ``request_id`` takes precedence so queued clinical files cannot be skipped.
    """

    clean_request_id = str(request_id or "").strip()
    if clean_request_id:
        return _process_queued_documents(clean_request_id, prompt)
    if not str(input_value or "").strip():
        return "Provide input_value for one document or request_id for queued documents."
    return _extract_document(input_value, prompt)


async def upload_file(request):
    from starlette.responses import JSONResponse

    try:
        _cleanup_stale_queues()
        content_type = request.headers.get("content-type", "")

        if content_type.startswith("multipart/form-data"):
            form = await request.form()
            request_id = str(form.get("request_id") or "")
            source = str(form.get("source") or "case")
            uploaded = form.get("file")
            if uploaded is None or not hasattr(uploaded, "read"):
                raise ValueError("Missing file.")
            content = await uploaded.read()
            filename = getattr(uploaded, "filename", None) or str(
                form.get("filename") or "document"
            )
            mime_type = getattr(uploaded, "content_type", None)
        else:
            data = await request.json()
            request_id = str(data.get("request_id") or "")
            source = str(data.get("source") or "case")
            filename = str(data.get("filename") or "document")
            encoded = str(data.get("content") or "")
            if not encoded:
                raise ValueError("Missing file content.")
            try:
                content = base64.b64decode(encoded, validate=True)
            except ValueError as error:
                raise ValueError("Invalid base64 file content.") from error
            mime_type = data.get("content_type")

        entry = _store_queued_file(
            request_id,
            filename,
            content,
            source,
            mime_type,
        )
        return JSONResponse(
            {
                "status": "queued",
                "request_id": _safe_request_id(request_id),
                "file": {
                    "name": entry["original_name"],
                    "source": entry["source"],
                    "size": entry["size"],
                },
            },
            status_code=201,
        )
    except ValueError as error:
        return JSONResponse({"status": "error", "message": str(error)}, status_code=400)
    except Exception as error:
        logger.exception("Clinical file upload failed")
        return JSONResponse(
            {"status": "error", "message": "The document could not be queued."},
            status_code=500,
        )


async def queue_status(request):
    from starlette.responses import JSONResponse

    try:
        request_id = _safe_request_id(request.path_params["request_id"])
        entries = _read_manifest(_queue_dir(request_id))
        return JSONResponse(
            {
                "request_id": request_id,
                "count": len(entries),
                "files": [
                    {
                        "name": entry.get("original_name"),
                        "source": entry.get("source"),
                        "size": entry.get("size"),
                    }
                    for entry in entries
                ],
            }
        )
    except ValueError as error:
        return JSONResponse({"status": "error", "message": str(error)}, status_code=400)


async def discard_queue(request):
    from starlette.responses import JSONResponse

    try:
        request_id = _safe_request_id(request.path_params["request_id"])
        shutil.rmtree(_queue_dir(request_id), ignore_errors=True)
        logger.info("Discarded clinical queue %s", request_id)
        return JSONResponse({"status": "discarded", "request_id": request_id})
    except ValueError as error:
        return JSONResponse({"status": "error", "message": str(error)}, status_code=400)


async def discard_all_queues(_request):
    from starlette.responses import JSONResponse

    cleared = _clear_all_queues()
    return JSONResponse({"status": "cleared", "cleared": cleared})


async def health(_request):
    from starlette.responses import JSONResponse

    _cleanup_stale_queues()
    queue_count = (
        sum(1 for item in TEMP_HOLD_ROOT.iterdir() if item.is_dir())
        if TEMP_HOLD_ROOT.exists()
        else 0
    )
    return JSONResponse(
        {
            "status": "ok",
            "service": "msob-clinical-documents",
            "queued_requests": queue_count,
            "vercel_local_access": True,
            "vercel_private_network": "cors-v1",
        }
    )


def build_app():
    from starlette.datastructures import Headers
    from starlette.middleware.cors import CORSMiddleware
    from starlette.responses import JSONResponse

    app = mcp.sse_app()

    class ProtectLocalQueueRoutes:
        def __init__(self, inner_app):
            self.inner_app = inner_app

        async def __call__(self, scope, receive, send):
            if scope["type"] == "http":
                path = str(scope.get("path") or "")
                if path == "/upload" or path == "/queue" or path.startswith("/queue/"):
                    headers = Headers(scope=scope)
                    forwarded_request = any(
                        headers.get(name)
                        for name in (
                            "forwarded",
                            "x-forwarded-for",
                            "x-forwarded-host",
                            "x-forwarded-proto",
                        )
                    )
                    host = (headers.get("host") or "").split(",", 1)[0].strip().lower()
                    local_host = (
                        host == "localhost"
                        or host.startswith("localhost:")
                        or host == "127.0.0.1"
                        or host.startswith("127.0.0.1:")
                        or host == "[::1]"
                        or host.startswith("[::1]:")
                    )
                    if forwarded_request or not local_host:
                        response = JSONResponse(
                            {
                                "status": "error",
                                "message": (
                                    "Clinical queue routes are available only "
                                    "from localhost."
                                ),
                            },
                            status_code=403,
                        )
                        await response(scope, receive, send)
                        return
            await self.inner_app(scope, receive, send)

    app.add_middleware(ProtectLocalQueueRoutes)
    app.add_middleware(
        CORSMiddleware,
        allow_origins=BROWSER_ORIGINS,
        allow_credentials=False,
        allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
        allow_headers=["*"],
        allow_private_network=True,
    )
    app.add_route("/health", health, methods=["GET"])
    app.add_route("/upload", upload_file, methods=["POST"])
    app.add_route("/queue", discard_all_queues, methods=["DELETE"])
    app.add_route("/queue/{request_id}", queue_status, methods=["GET"])
    app.add_route("/queue/{request_id}", discard_queue, methods=["DELETE"])
    return app


if __name__ == "__main__":
    import uvicorn

    TEMP_HOLD_ROOT.mkdir(parents=True, exist_ok=True)
    _cleanup_stale_queues()
    logger.info("Starting MSOB clinical-document MCP server on port %s", mcp.settings.port)
    uvicorn.run(
        build_app(),
        host=os.getenv("MCP_HOST", "0.0.0.0"),
        port=int(os.getenv("MCP_PORT", "8002")),
    )
